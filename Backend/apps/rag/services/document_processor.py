# """
# Document Processing Service
# Handles file upload, text extraction, chunking, and vectorization
# """
# import logging
# from typing import Dict, Any, List
# import PyPDF2
# import docx
# import io
# from django.core.files.uploadedfile import UploadedFile
# from django.conf import settings

# logger = logging.getLogger(__name__)


# class DocumentProcessor:
#     """
#     Processes uploaded documents for RAG system.
    
#     Handles:
#     - Text extraction (PDF, DOCX, TXT)
#     - Text chunking
#     - Embedding generation
#     - Vector store insertion
#     """
    
#     def __init__(self, vector_store, embedding_service, chunk_size=None, chunk_overlap=None):
#         """
#         Initialize document processor.
        
#         Args:
#             vector_store: Vector store instance
#             embedding_service: Embedding service instance
#             chunk_size: Size of text chunks
#             chunk_overlap: Overlap between chunks
#         """
#         self.vector_store = vector_store
#         self.embedding_service = embedding_service
#         self.chunk_size = chunk_size or getattr(settings, 'CHUNK_SIZE', 800)
#         self.chunk_overlap = chunk_overlap or getattr(settings, 'CHUNK_OVERLAP', 100)
    
#     async def process_document(self, file: UploadedFile, document_id: str) -> Dict[str, Any]:
#         """
#         Process uploaded document.
        
#         Args:
#             file: Uploaded file
#             document_id: Document UUID
            
#         Returns:
#             Processing result dictionary
#         """
#         # Extract text
#         text = self._extract_text(file)
        
#         if not text or len(text.strip()) < 10:
#             raise ValueError("No extractable text found in document")
        
#         logger.info(f"[Processor] Extracted {len(text)} characters")
        
#         # Chunk text
#         chunks = self._chunk_text(text)
        
#         if not chunks:
#             raise ValueError("No chunks generated from document")
        
#         logger.info(f"[Processor] Created {len(chunks)} chunks")
        
#         # Generate embeddings
#         embeddings = self.embedding_service.embed_texts(chunks)
        
#         # Prepare metadata
#         metadatas = [
#             {
#                 "source": file.name,
#                 "content_type": file.content_type or 'application/octet-stream',
#                 "chunk_index": i,
#                 "document_id": document_id,
#                 "chunk_size": len(chunk)
#             }
#             for i, chunk in enumerate(chunks)
#         ]
        
#         # Generate IDs
#         ids = [f"{document_id}_chunk_{i}" for i in range(len(chunks))]
        
#         # Add to vector store
#         self.vector_store.add_documents(
#             documents=chunks,
#             embeddings=embeddings,
#             metadata=metadatas,
#             ids=ids
#         )
        
#         logger.info(f"[Processor] Added {len(chunks)} chunks to vector store")
        
#         return {
#             "chunks_created": len(chunks),
#             "text_length": len(text),
#             "document_id": document_id
#         }
    
#     def _extract_text(self, file: UploadedFile) -> str:
#         """Extract text from uploaded file"""
#         content = file.read()
        
#         if file.content_type == 'application/pdf' or file.name.endswith('.pdf'):
#             return self._extract_pdf(content)
#         elif file.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or file.name.endswith('.docx'):
#             return self._extract_docx(content)
#         elif file.content_type == 'text/plain' or file.name.endswith('.txt'):
#             return content.decode('utf-8', errors='ignore')
#         else:
#             raise ValueError(f"Unsupported file type: {file.content_type}")
    
#     def _extract_pdf(self, content: bytes) -> str:
#         """Extract text from PDF"""
#         try:
#             pdf_file = io.BytesIO(content)
#             pdf_reader = PyPDF2.PdfReader(pdf_file)
#             text = ""
#             for page in pdf_reader.pages:
#                 text += page.extract_text() + "\n"
#             return text.strip()
#         except Exception as e:
#             logger.error(f"PDF extraction failed: {e}")
#             raise ValueError(f"Failed to extract PDF: {e}")
    
#     def _extract_docx(self, content: bytes) -> str:
#         """Extract text from DOCX"""
#         try:
#             doc_file = io.BytesIO(content)
#             doc = docx.Document(doc_file)
#             text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
#             return text.strip()
#         except Exception as e:
#             logger.error(f"DOCX extraction failed: {e}")
#             raise ValueError(f"Failed to extract DOCX: {e}")
    
#     def _chunk_text(self, text: str) -> List[str]:
#         """
#         Split text into overlapping chunks.
        
#         Args:
#             text: Input text
            
#         Returns:
#             List of text chunks
#         """
#         if not text:
#             return []
        
#         words = text.split()
#         if len(words) <= self.chunk_size:
#             return [" ".join(words)]
        
#         chunks = []
#         start = 0
        
#         while start < len(words):
#             end = start + self.chunk_size
#             chunk = " ".join(words[start:end])
            
#             if chunk.strip():
#                 chunks.append(chunk.strip())
            
#             if end >= len(words):
#                 break
            
#             start = end - self.chunk_overlap
        
#         return chunks












# """
# Document Processing Service
# Handles file upload, text extraction, chunking, and vectorization
# Supports: PDF, DOCX, TXT, CSV
# Tables: comma CSV, pipe-delimited, tab TSV, semicolon-delimited, PDF tables, DOCX tables
# """
# import logging
# import csv as csv_module
# import io
# from typing import Dict, Any, List
# from django.core.files.uploadedfile import UploadedFile
# from django.conf import settings

# logger = logging.getLogger(__name__)


# class DocumentProcessor:
#     """
#     Processes uploaded documents for RAG system.

#     Handles:
#     - Text extraction (PDF, DOCX, TXT, CSV)
#     - Table extraction (CSV, pipe-delimited, PDF tables, DOCX tables)
#     - Text chunking
#     - Embedding generation
#     - Vector store insertion
#     """

#     def __init__(self, vector_store, embedding_service, chunk_size=None, chunk_overlap=None):
#         """
#         Initialize document processor.

#         Args:
#             vector_store:      Vector store instance
#             embedding_service: Embedding service instance
#             chunk_size:        Size of text chunks
#             chunk_overlap:     Overlap between chunks
#         """
#         self.vector_store      = vector_store
#         self.embedding_service = embedding_service
#         self.chunk_size        = chunk_size    or getattr(settings, 'CHUNK_SIZE',    800)
#         self.chunk_overlap     = chunk_overlap or getattr(settings, 'CHUNK_OVERLAP', 100)

#     # =========================================================
#     #  MAIN ENTRY POINT
#     # =========================================================
#     async def process_document(self, file: UploadedFile, document_id: str) -> Dict[str, Any]:
#         """
#         Process uploaded document.

#         Args:
#             file:        Uploaded file (PDF / DOCX / TXT / CSV)
#             document_id: Document UUID

#         Returns:
#             Processing result dictionary
#         """
#         # Extract text (plain text + tables converted to natural language)
#         text = self._extract_text(file)

#         if not text or len(text.strip()) < 10:
#             raise ValueError("No extractable text found in document")

#         logger.info(f"[Processor] Extracted {len(text)} characters from {file.name}")

#         # Chunk text
#         chunks = self._chunk_text(text)

#         if not chunks:
#             raise ValueError("No chunks generated from document")

#         logger.info(f"[Processor] Created {len(chunks)} chunks")

#         # Generate embeddings (batch call — same as original)
#         embeddings = self.embedding_service.embed_texts(chunks)

#         # Prepare metadata
#         metadatas = [
#             {
#                 "source":       file.name,
#                 "content_type": file.content_type or 'application/octet-stream',
#                 "chunk_index":  i,
#                 "document_id":  document_id,
#                 "chunk_size":   len(chunk)
#             }
#             for i, chunk in enumerate(chunks)
#         ]

#         # Generate IDs
#         ids = [f"{document_id}_chunk_{i}" for i in range(len(chunks))]

#         # Add to vector store (batch call — same as original)
#         self.vector_store.add_documents(
#             documents  = chunks,
#             embeddings = embeddings,
#             metadata   = metadatas,
#             ids        = ids
#         )

#         logger.info(f"[Processor] Added {len(chunks)} chunks to vector store")

#         return {
#             "chunks_created": len(chunks),
#             "text_length":    len(text),
#             "document_id":    document_id
#         }

#     # =========================================================
#     #  TEXT EXTRACTION — routes by file type
#     # =========================================================
#     def _extract_text(self, file: UploadedFile) -> str:
#         """Extract text from uploaded file based on its type."""
#         content     = file.read()
#         name_lower  = file.name.lower()
#         ct          = file.content_type or ''

#         if name_lower.endswith('.pdf') or 'pdf' in ct:
#             return self._extract_pdf(content)

#         elif name_lower.endswith('.docx') or 'wordprocessingml' in ct:
#             return self._extract_docx(content)

#         elif name_lower.endswith('.csv') or 'csv' in ct:
#             # CSV always treated as a table file
#             raw = self._decode(content)
#             return self._parse_csv(raw, delimiter=',')

#         elif name_lower.endswith('.txt') or 'text/plain' in ct:
#             return self._extract_txt(content)

#         else:
#             raise ValueError(
#                 f"Unsupported file type: {file.content_type or file.name}. "
#                 "Allowed: PDF, DOCX, TXT, CSV"
#             )

#     # =========================================================
#     #  PDF — plain text + embedded tables via pdfplumber
#     # =========================================================
#     def _extract_pdf(self, content: bytes) -> str:
#         """
#         Extract text and tables from PDF.
#         Uses pdfplumber for better table support.
#         Falls back to PyPDF2 if pdfplumber is unavailable.
#         """
#         # --- Try pdfplumber first (best table support) ---
#         try:
#             import pdfplumber
#             parts = []

#             with pdfplumber.open(io.BytesIO(content)) as pdf:
#                 for page_num, page in enumerate(pdf.pages, 1):

#                     # Tables on this page
#                     try:
#                         for raw_table in (page.extract_tables() or []):
#                             if raw_table and len(raw_table) >= 2:
#                                 table_text = self._table_rows_to_text(raw_table)
#                                 if table_text:
#                                     parts.append(table_text)
#                     except Exception as e:
#                         logger.warning(f"PDF table extract p{page_num}: {e}")

#                     # Plain text on this page
#                     try:
#                         page_text = page.extract_text()
#                         if page_text and page_text.strip():
#                             parts.append(page_text.strip())
#                     except Exception as e:
#                         logger.warning(f"PDF text extract p{page_num}: {e}")

#             result = "\n\n".join(parts).strip()
#             if result:
#                 return result
#             # Fall through to PyPDF2 if pdfplumber got nothing

#         except ImportError:
#             logger.warning("[Processor] pdfplumber not found, falling back to PyPDF2")

#         # --- Fallback: PyPDF2 (no table support but always available) ---
#         try:
#             import PyPDF2
#             pdf_file   = io.BytesIO(content)
#             pdf_reader = PyPDF2.PdfReader(pdf_file)
#             text = ""
#             for page in pdf_reader.pages:
#                 text += (page.extract_text() or "") + "\n"
#             return text.strip()
#         except Exception as e:
#             logger.error(f"PDF extraction failed: {e}")
#             raise ValueError(f"Failed to extract PDF: {e}")

#     # =========================================================
#     #  DOCX — paragraphs + tables in document order
#     # =========================================================
#     def _extract_docx(self, content: bytes) -> str:
#         """Extract paragraphs and tables from DOCX in document order."""
#         try:
#             import docx as docx_lib
#             doc_file = io.BytesIO(content)
#             doc      = docx_lib.Document(doc_file)
#             parts    = []

#             WNS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

#             # Walk body elements in order so text and tables are interleaved correctly
#             for element in doc.element.body:
#                 tag = element.tag.split('}')[-1]

#                 if tag == 'p':
#                     # Regular paragraph
#                     text = ''.join(
#                         n.text or '' for n in element.iter()
#                         if n.tag.endswith('}t')
#                     )
#                     if text.strip():
#                         parts.append(text.strip())

#                 elif tag == 'tbl':
#                     # Table — extract rows and cells
#                     rows = []
#                     for tr in element.findall(f'.//{WNS}tr'):
#                         cells = []
#                         for tc in tr.findall(f'.//{WNS}tc'):
#                             cell_text = ''.join(
#                                 n.text or '' for n in tc.iter()
#                                 if n.tag.endswith('}t')
#                             ).strip()
#                             cells.append(cell_text)
#                         if cells:
#                             rows.append(cells)

#                     if rows:
#                         table_text = self._table_rows_to_text(rows)
#                         if table_text:
#                             parts.append(table_text)

#             return "\n\n".join(parts).strip()

#         except Exception as e:
#             logger.error(f"DOCX extraction failed: {e}")
#             raise ValueError(f"Failed to extract DOCX: {e}")

#     # =========================================================
#     #  TXT — auto-detects CSV / pipe / TSV / plain
#     # =========================================================
#     def _extract_txt(self, content: bytes) -> str:
#         """
#         Extract text from TXT file.
#         Auto-detects table formats: CSV, TSV, pipe-delimited, semicolon.
#         """
#         raw   = self._decode(content)
#         lines = [l for l in raw.splitlines() if l.strip()]

#         if not lines:
#             return ""

#         # Detection order: comma -> tab -> semicolon -> pipe -> plain text
#         if self._looks_like_delimited(lines, ','):
#             logger.info("[Processor] TXT: detected CSV (comma)")
#             return self._parse_csv(raw, delimiter=',')

#         if self._looks_like_delimited(lines, '\t'):
#             logger.info("[Processor] TXT: detected TSV (tab)")
#             return self._parse_csv(raw, delimiter='\t')

#         if self._looks_like_delimited(lines, ';'):
#             logger.info("[Processor] TXT: detected semicolon-delimited")
#             return self._parse_csv(raw, delimiter=';')

#         if self._looks_like_delimited(lines, '|'):
#             logger.info("[Processor] TXT: detected pipe-delimited table")
#             return self._parse_pipe(lines)

#         # Plain text
#         logger.info("[Processor] TXT: plain text mode")
#         return raw.strip()

#     # =========================================================
#     #  TABLE FORMAT DETECTION
#     # =========================================================
#     def _looks_like_delimited(self, lines: List[str], delimiter: str) -> bool:
#         """
#         Returns True if the first few lines have a consistent
#         count of the given delimiter (>= 1 per line).
#         """
#         sample = lines[:min(5, len(lines))]
#         counts = [line.count(delimiter) for line in sample]
#         return len(set(counts)) == 1 and counts[0] >= 1

#     # =========================================================
#     #  CSV / TSV / SEMICOLON  ->  natural language rows
#     # =========================================================
#     def _parse_csv(self, raw: str, delimiter: str = ',') -> str:
#         """
#         Converts every CSV data row into a natural-language string.

#         Input:
#             Student_ID, First_Name, GPA
#             S001,       Emma,       3.8

#         Output:
#             Student_ID: S001 | First_Name: Emma | GPA: 3.8
#         """
#         reader = csv_module.reader(io.StringIO(raw), delimiter=delimiter)
#         rows   = [r for r in reader if any(c.strip() for c in r)]

#         if len(rows) < 2:
#             # No header or only one row — return raw text
#             return raw.strip()

#         headers = [h.strip() for h in rows[0]]
#         result  = []

#         for row in rows[1:]:
#             pairs = []
#             for header, cell in zip(headers, row):
#                 cell = cell.strip()
#                 if cell:
#                     pairs.append(f"{header}: {cell}")
#             if pairs:
#                 result.append(" | ".join(pairs))

#         logger.info(f"[Processor] CSV: converted {len(result)} rows to natural language")
#         return "\n".join(result)

#     # =========================================================
#     #  PIPE-DELIMITED  ->  natural language rows
#     # =========================================================
#     def _parse_pipe(self, lines: List[str]) -> str:
#         """
#         Parses markdown-style pipe tables.

#         Input:
#             | ID  | Name  | Score |
#             |-----|-------|-------|
#             | 101 | Ahmed | 95    |

#         Output:
#             ID: 101 | Name: Ahmed | Score: 95
#         """
#         # Remove separator lines like |---|---|
#         clean = [
#             line for line in lines
#             if '|' in line and not all(c in '-|:+ \t' for c in line)
#         ]
#         if not clean:
#             return "\n".join(lines)

#         rows = []
#         for line in clean:
#             cells = [c.strip() for c in line.split('|')]
#             cells = [c for c in cells if c]  # drop empty edge cells
#             if cells:
#                 rows.append(cells)

#         return self._table_rows_to_text(rows)

#     # =========================================================
#     #  GENERIC TABLE ROWS  ->  natural language (PDF + DOCX + pipe)
#     # =========================================================
#     def _table_rows_to_text(self, rows: List[List]) -> str:
#         """
#         Converts any list-of-lists table (first row = headers) into
#         one natural-language string per data row.

#         [['Name', 'Age'], ['Ahmed', '25'], ['Sara', '22']]
#         ->
#         Name: Ahmed | Age: 25
#         Name: Sara  | Age: 22
#         """
#         if not rows or len(rows) < 2:
#             return ""

#         # Clean every cell
#         cleaned = [
#             [str(c).strip() if c is not None else "" for c in row]
#             for row in rows
#         ]

#         headers   = cleaned[0]
#         data_rows = cleaned[1:]

#         # Skip separator rows (all dashes/equals)
#         data_rows = [
#             row for row in data_rows
#             if any(cell for cell in row)
#             and not all(set(cell) <= {'-', '=', '_', ' ', ''} for cell in row)
#         ]

#         if not data_rows:
#             return ""

#         result = []
#         for row in data_rows:
#             pairs = []
#             for i, (header, cell) in enumerate(zip(headers, row)):
#                 if cell:
#                     label = header if header else f"col_{i}"
#                     pairs.append(f"{label}: {cell}")
#             if pairs:
#                 result.append(" | ".join(pairs))

#         return "\n".join(result)

#     # =========================================================
#     #  CHUNKING — line-aware so table rows are never split
#     # =========================================================
#     def _chunk_text(self, text: str) -> List[str]:
#         """
#         Split text into overlapping chunks.
#         Splits on newlines first so individual table rows stay intact.
#         """
#         if not text:
#             return []

#         words = text.split()

#         # Short document — return as single chunk
#         if len(words) <= self.chunk_size:
#             return [" ".join(words)]

#         chunks = []
#         start  = 0

#         while start < len(words):
#             end   = start + self.chunk_size
#             chunk = " ".join(words[start:end])

#             if chunk.strip():
#                 chunks.append(chunk.strip())

#             if end >= len(words):
#                 break

#             start = end - self.chunk_overlap

#         return chunks

#     # =========================================================
#     #  DECODE HELPER — tries multiple encodings
#     # =========================================================
#     def _decode(self, file_bytes: bytes) -> str:
#         """Try common encodings in order, never crash on bad bytes."""
#         for enc in ('utf-8', 'utf-8-sig', 'latin-1', 'cp1252'):
#             try:
#                 return file_bytes.decode(enc)
#             except UnicodeDecodeError:
#                 continue
#         return file_bytes.decode('utf-8', errors='replace')
    





# """
# ═══════════════════════════════════════════════════════════════════════════════
#         GENERIC RAG FILE PROCESSOR — CSV, TXT, PDF, DOCX
#         Works with ANY file — no hardcoded column names
#         Production-grade approach used by LangChain / LlamaIndex / Haystack
# ═══════════════════════════════════════════════════════════════════════════════
# """

# import os
# import re
# import csv
# import json
# import hashlib
# import pandas as pd
# from pathlib import Path
# from typing import List, Dict, Any, Optional, Tuple
# from dataclasses import dataclass, field


# # ─────────────────────────────────────────────────────────────────────────────
# #  DATA STRUCTURES
# # ─────────────────────────────────────────────────────────────────────────────

# @dataclass
# class Document:
#     """A single indexable unit for ChromaDB / any vector store."""
#     text: str                          # Natural language text to embed
#     metadata: Dict[str, Any]           # Filterable fields
#     doc_id: str                        # Unique ID
#     source_file: str                   # Original filename
#     chunk_index: int = 0              # Position within file
#     row_index: Optional[int] = None   # For tabular data: which row


# @dataclass
# class ProcessingResult:
#     """Result returned after processing any file."""
#     documents: List[Document]
#     file_type: str
#     total_rows: int
#     columns_detected: List[str]
#     warnings: List[str] = field(default_factory=list)


# # ─────────────────────────────────────────────────────────────────────────────
# #  STEP 1 — AUTO-DETECT FILE TYPE & STRUCTURE
# # ─────────────────────────────────────────────────────────────────────────────

# class FileTypeDetector:
#     """
#     Automatically detects file type and internal structure.
#     No hardcoding — works with any file.
#     """

#     TABULAR_EXTENSIONS = {'.csv', '.tsv', '.txt'}
#     TEXT_EXTENSIONS = {'.txt', '.md', '.log'}
#     DOC_EXTENSIONS = {'.pdf', '.docx', '.doc'}

#     @staticmethod
#     def detect(file_path: str) -> Dict[str, Any]:
#         path = Path(file_path)
#         ext = path.suffix.lower()
        
#         result = {
#             "extension": ext,
#             "file_type": "unknown",
#             "is_tabular": False,
#             "delimiter": None,
#             "has_header": False,
#             "encoding": "utf-8",
#         }

#         # ── Detect encoding safely ──
#         result["encoding"] = FileTypeDetector._detect_encoding(file_path)

#         # ── CSV — always tabular ──
#         if ext == ".csv":
#             result["file_type"] = "csv"
#             result["is_tabular"] = True
#             result["delimiter"] = ","
#             result["has_header"] = True
#             return result

#         # ── TSV ──
#         if ext == ".tsv":
#             result["file_type"] = "tsv"
#             result["is_tabular"] = True
#             result["delimiter"] = "\t"
#             result["has_header"] = True
#             return result

#         # ── TXT — could be tabular OR free text ──
#         if ext == ".txt":
#             tabular_info = FileTypeDetector._analyze_txt(file_path, result["encoding"])
#             result.update(tabular_info)
#             return result

#         # ── PDF / DOCX ──
#         if ext in FileTypeDetector.DOC_EXTENSIONS:
#             result["file_type"] = ext.strip(".")
#             result["is_tabular"] = False
#             return result

#         return result

#     @staticmethod
#     def _detect_encoding(file_path: str) -> str:
#         """Try common encodings — never crash on encoding."""
#         for enc in ["utf-8", "latin-1", "utf-16", "cp1252"]:
#             try:
#                 with open(file_path, "r", encoding=enc) as f:
#                     f.read(1024)
#                 return enc
#             except (UnicodeDecodeError, Exception):
#                 continue
#         return "latin-1"  # fallback

#     @staticmethod
#     def _analyze_txt(file_path: str, encoding: str) -> Dict[str, Any]:
#         """
#         Analyze a .txt file to determine if it's tabular or free text.
#         Detects delimiter automatically — comma, tab, pipe, semicolon.
#         """
#         try:
#             with open(file_path, "r", encoding=encoding) as f:
#                 sample_lines = [f.readline() for _ in range(10)]
#             sample_lines = [l.strip() for l in sample_lines if l.strip()]
#         except Exception:
#             return {"file_type": "txt_freetext", "is_tabular": False}

#         if not sample_lines:
#             return {"file_type": "txt_freetext", "is_tabular": False}

#         # Count delimiters per line
#         delimiters = {",": 0, "\t": 0, "|": 0, ";": 0}
#         for line in sample_lines:
#             for delim in delimiters:
#                 delimiters[delim] += line.count(delim)

#         # Pick the most common delimiter
#         best_delim = max(delimiters, key=delimiters.get)
#         best_count = delimiters[best_delim]

#         # If average occurrences per line > 1, likely tabular
#         avg_per_line = best_count / max(len(sample_lines), 1)
#         if avg_per_line >= 1:
#             # Check consistency — each line should have same number of fields
#             field_counts = [len(line.split(best_delim)) for line in sample_lines]
#             is_consistent = len(set(field_counts)) <= 2  # allow 1-2 variation

#             if is_consistent:
#                 # Try to detect header — first row has mostly text, rest have numbers
#                 first_row_fields = sample_lines[0].split(best_delim)
#                 has_header = all(
#                     not re.match(r"^[0-9.]+$", f.strip()) 
#                     for f in first_row_fields[:5]
#                 )
#                 return {
#                     "file_type": "txt_tabular",
#                     "is_tabular": True,
#                     "delimiter": best_delim,
#                     "has_header": has_header,
#                 }

#         return {"file_type": "txt_freetext", "is_tabular": False}


# # ─────────────────────────────────────────────────────────────────────────────
# #  STEP 2 — GENERIC TABULAR PROCESSOR (No hardcoded columns!)
# # ─────────────────────────────────────────────────────────────────────────────

# class TabularProcessor:
#     """
#     Converts ANY tabular file (CSV or tabular TXT) into RAG-ready documents.
    
#     KEY INSIGHT: Instead of hardcoding column names, we auto-detect them
#     and build "FieldName: Value" text for every row dynamically.
    
#     This is EXACTLY how production systems like LlamaIndex's CSVReader work.
#     """

#     def __init__(self, chunk_size: int = 50, overlap: int = 5):
#         self.chunk_size = chunk_size  # rows per chunk for large files
#         self.overlap = overlap

#     def process(self, file_path: str, file_info: Dict) -> ProcessingResult:
#         """Main entry — process any tabular file generically."""
#         warnings = []
        
#         # ── Load into DataFrame ──
#         try:
#             df = self._load_dataframe(file_path, file_info)
#         except Exception as e:
#             return ProcessingResult(
#                 documents=[], file_type=file_info["file_type"],
#                 total_rows=0, columns_detected=[],
#                 warnings=[f"Failed to load file: {str(e)}"]
#             )

#         if df.empty:
#             return ProcessingResult(
#                 documents=[], file_type=file_info["file_type"],
#                 total_rows=0, columns_detected=list(df.columns),
#                 warnings=["File is empty"]
#             )

#         # ── Clean column names ──
#         df.columns = [self._clean_col_name(c) for c in df.columns]
        
#         # ── Remove completely empty rows ──
#         df = df.dropna(how="all")
        
#         # ── Detect column types automatically ──
#         col_types = self._detect_column_types(df)
        
#         # ── Build documents ──
#         documents = []
#         filename = Path(file_path).name

#         for idx, row in df.iterrows():
#             doc = self._row_to_document(
#                 row=row,
#                 row_index=idx,
#                 col_types=col_types,
#                 filename=filename,
#                 total_rows=len(df)
#             )
#             documents.append(doc)

#         # ── Also create summary chunks for large files ──
#         if len(df) > 20:
#             summary_docs = self._create_summary_chunks(df, filename, col_types)
#             documents.extend(summary_docs)
            
#         if len(df) > 1000:
#             warnings.append(f"Large file: {len(df)} rows. Consider batch indexing.")

#         return ProcessingResult(
#             documents=documents,
#             file_type=file_info["file_type"],
#             total_rows=len(df),
#             columns_detected=list(df.columns),
#             warnings=warnings
#         )

#     def _load_dataframe(self, file_path: str, file_info: Dict) -> pd.DataFrame:
#         """Load any tabular file into a DataFrame — fully generic."""
#         enc = file_info.get("encoding", "utf-8")
#         delim = file_info.get("delimiter", ",")
        
#         if file_info["extension"] == ".csv":
#             # Let pandas auto-detect separator if needed
#             try:
#                 return pd.read_csv(file_path, encoding=enc, dtype=str)
#             except Exception:
#                 return pd.read_csv(file_path, encoding=enc, sep=None, 
#                                    engine="python", dtype=str)

#         elif file_info["file_type"] == "txt_tabular":
#             return pd.read_csv(
#                 file_path, 
#                 sep=delim,
#                 encoding=enc,
#                 dtype=str,
#                 header=0 if file_info.get("has_header") else None,
#                 on_bad_lines="skip"  # skip malformed rows gracefully
#             )
        
#         return pd.DataFrame()

#     @staticmethod
#     def _clean_col_name(col: str) -> str:
#         """Normalize column names for consistent output."""
#         col = str(col).strip()
#         col = re.sub(r"[_\-]+", " ", col)  # underscores → spaces
#         col = col.title()                   # Title Case
#         col = re.sub(r"\s+", " ", col)
#         return col

#     @staticmethod
#     def _detect_column_types(df: pd.DataFrame) -> Dict[str, str]:
#         """
#         Auto-detect what kind of data each column holds.
#         Used to write smarter natural language descriptions.
#         """
#         col_types = {}
#         for col in df.columns:
#             sample = df[col].dropna().head(10)
            
#             # Try numeric
#             numeric_count = sum(
#                 1 for v in sample 
#                 if re.match(r"^-?[0-9]+\.?[0-9]*$", str(v).strip())
#             )
#             if numeric_count >= len(sample) * 0.8:
#                 col_types[col] = "numeric"
#                 continue
            
#             # Try email
#             email_count = sum(1 for v in sample if "@" in str(v))
#             if email_count >= len(sample) * 0.5:
#                 col_types[col] = "email"
#                 continue

#             # Try year
#             year_count = sum(
#                 1 for v in sample 
#                 if re.match(r"^(19|20)\d{2}$", str(v).strip())
#             )
#             if year_count >= len(sample) * 0.8:
#                 col_types[col] = "year"
#                 continue

#             # Try percentage
#             pct_count = sum(
#                 1 for v in sample 
#                 if re.match(r"^[0-9]+\.?[0-9]*%?$", str(v).strip()) 
#                 and 0 <= float(re.sub(r"%", "", str(v))) <= 100
#             )
#             if pct_count >= len(sample) * 0.8:
#                 col_types[col] = "percentage"
#                 continue

#             # Try ID column
#             col_lower = col.lower()
#             if any(kw in col_lower for kw in ["id", "code", "number", "no"]):
#                 col_types[col] = "identifier"
#                 continue

#             col_types[col] = "text"
        
#         return col_types

#     def _row_to_document(
#         self, 
#         row: pd.Series, 
#         row_index: int,
#         col_types: Dict[str, str],
#         filename: str,
#         total_rows: int
#     ) -> Document:
#         """
#         Convert ONE row into a natural language Document.
        
#         GENERIC — works with any columns, any file.
        
#         Example output for Student CSV:
#           "Student record from Students_Data.csv:
#            Student Id: S024 | First Name: Daniel | Last Name: Clark |
#            Age: 16 | Gpa: 3.5 | Math Score: 85 ..."
           
#         Example output for a Sales CSV with different columns:
#           "Sales record from sales_q1.csv:
#            Product: Widget A | Region: North | Revenue: 50000 ..."
#         """
#         # ── Infer a "record type" from the filename ──
#         record_type = self._infer_record_type(filename)
        
#         # ── Build field descriptions ──
#         parts = []
#         metadata = {
#             "source_file": filename,
#             "row_index": row_index,
#             "record_type": record_type,
#         }
        
#         for col in row.index:
#             raw_val = row[col]
#             if pd.isna(raw_val) or str(raw_val).strip() in ("", "nan", "None"):
#                 continue
            
#             val = str(raw_val).strip()
#             col_type = col_types.get(col, "text")
            
#             # Format value based on type
#             if col_type == "percentage":
#                 display_val = f"{val}%"
#             elif col_type == "email":
#                 display_val = val
#             else:
#                 display_val = val

#             parts.append(f"{col}: {display_val}")
            
#             # Add to metadata for filtering (store original value)
#             # Key: lowercase, no spaces — for ChromaDB where clause
#             meta_key = col.lower().replace(" ", "_")
#             metadata[meta_key] = val

#         # ── Build the final text ──
#         fields_text = " | ".join(parts)
#         text = f"{record_type} from {filename}:\n{fields_text}"
        
#         # ── Also add a semantic-friendly summary line ──
#         name_val = self._extract_name(row, col_types)
#         if name_val:
#             text = f"{name_val} — {record_type} from {filename}:\n{fields_text}"

#         # ── Generate stable unique ID ──
#         doc_id = hashlib.md5(f"{filename}_{row_index}".encode()).hexdigest()[:16]

#         return Document(
#             text=text,
#             metadata=metadata,
#             doc_id=doc_id,
#             source_file=filename,
#             row_index=row_index
#         )

#     @staticmethod
#     def _infer_record_type(filename: str) -> str:
#         """
#         Infer what kind of records the file contains from its name.
#         Generic fallback: "Record"
        
#         students_data.csv → "Student record"
#         sales_report.csv  → "Sales report record"
#         employees.txt     → "Employee record"
#         """
#         stem = Path(filename).stem.lower()
#         stem = re.sub(r"[_\-\.]", " ", stem)
        
#         # Remove common suffixes/numbers
#         stem = re.sub(r"\d+", "", stem).strip()
#         stem = re.sub(r"\b(data|file|report|sheet|export|list)\b", "", stem).strip()
#         stem = re.sub(r"\s+", " ", stem).strip()

#         if not stem:
#             return "Record"

#         # Singularize simple plurals
#         if stem.endswith("s") and len(stem) > 4:
#             stem = stem[:-1]

#         return stem.title() + " record"

#     @staticmethod
#     def _extract_name(row: pd.Series, col_types: Dict[str, str]) -> Optional[str]:
#         """
#         Try to find a 'name' in the row for a friendlier text prefix.
#         Looks for columns named: name, full_name, first_name+last_name, title, etc.
#         """
#         row_lower = {k.lower(): v for k, v in row.items()}
        
#         # Try: "name" column
#         for key in ["name", "full name", "fullname", "title"]:
#             if key in row_lower and str(row_lower[key]).strip():
#                 return str(row_lower[key]).strip()
        
#         # Try: first_name + last_name
#         first = row_lower.get("first name", row_lower.get("firstname", ""))
#         last = row_lower.get("last name", row_lower.get("lastname", ""))
#         if first and last:
#             return f"{str(first).strip()} {str(last).strip()}"
        
#         return None

#     def _create_summary_chunks(
#         self, 
#         df: pd.DataFrame, 
#         filename: str,
#         col_types: Dict[str, str]
#     ) -> List[Document]:
#         """
#         For large files, create summary chunks:
#         - Overall file summary (column names, row count, ranges)
#         - Numeric column statistics
        
#         These help answer aggregate questions like:
#         "How many students are in grade 11?"
#         "What is the average GPA?"
#         """
#         docs = []
#         record_type = self._infer_record_type(filename)

#         # ── File overview ──
#         col_summary = ", ".join(df.columns.tolist())
#         overview_text = (
#             f"Overview of {filename}: "
#             f"This file contains {len(df)} {record_type}s. "
#             f"Columns available: {col_summary}. "
#         )
        
#         # Add numeric summaries
#         numeric_cols = [c for c, t in col_types.items() if t == "numeric"]
#         for col in numeric_cols[:8]:  # limit to 8 columns
#             try:
#                 vals = pd.to_numeric(df[col], errors="coerce").dropna()
#                 if len(vals) > 0:
#                     overview_text += (
#                         f"{col} ranges from {vals.min():.2f} to {vals.max():.2f} "
#                         f"(average: {vals.mean():.2f}). "
#                     )
#             except Exception:
#                 pass

#         docs.append(Document(
#             text=overview_text,
#             metadata={"source_file": filename, "chunk_type": "summary"},
#             doc_id=hashlib.md5(f"{filename}_summary".encode()).hexdigest()[:16],
#             source_file=filename
#         ))

#         # ── Categorical summaries (e.g., unique cities, grade levels) ──
#         text_cols = [c for c, t in col_types.items() 
#                      if t == "text" and df[c].nunique() < 30]
#         for col in text_cols[:4]:
#             unique_vals = df[col].dropna().unique().tolist()
#             vals_str = ", ".join(str(v) for v in unique_vals[:20])
#             cat_text = (
#                 f"Unique values for '{col}' in {filename}: {vals_str}. "
#                 f"Total unique: {len(unique_vals)}."
#             )
#             docs.append(Document(
#                 text=cat_text,
#                 metadata={"source_file": filename, "chunk_type": "categorical_summary", "column": col},
#                 doc_id=hashlib.md5(f"{filename}_{col}_cat".encode()).hexdigest()[:16],
#                 source_file=filename
#             ))

#         return docs


# # ─────────────────────────────────────────────────────────────────────────────
# #  STEP 3 — FREE TEXT PROCESSOR (for non-tabular TXT, PDF, etc.)
# # ─────────────────────────────────────────────────────────────────────────────

# class FreeTextProcessor:
#     """
#     Handles plain text files that are NOT tabular.
#     Uses sliding window chunking — same as LangChain's RecursiveTextSplitter.
#     """

#     def __init__(self, chunk_size: int = 500, overlap: int = 50):
#         self.chunk_size = chunk_size  # characters per chunk
#         self.overlap = overlap

#     def process(self, file_path: str, file_info: Dict) -> ProcessingResult:
#         enc = file_info.get("encoding", "utf-8")
#         filename = Path(file_path).name

#         try:
#             with open(file_path, "r", encoding=enc) as f:
#                 text = f.read()
#         except Exception as e:
#             return ProcessingResult(
#                 documents=[], file_type="txt_freetext",
#                 total_rows=0, columns_detected=[],
#                 warnings=[f"Could not read file: {e}"]
#             )

#         # Split into chunks with overlap
#         chunks = self._sliding_window_split(text)
#         documents = []

#         for i, chunk_text in enumerate(chunks):
#             doc_id = hashlib.md5(f"{filename}_{i}".encode()).hexdigest()[:16]
#             documents.append(Document(
#                 text=chunk_text.strip(),
#                 metadata={
#                     "source_file": filename,
#                     "chunk_index": i,
#                     "chunk_type": "text_chunk"
#                 },
#                 doc_id=doc_id,
#                 source_file=filename,
#                 chunk_index=i
#             ))

#         return ProcessingResult(
#             documents=documents,
#             file_type="txt_freetext",
#             total_rows=len(chunks),
#             columns_detected=[]
#         )

#     def _sliding_window_split(self, text: str) -> List[str]:
#         """Split text into overlapping chunks."""
#         chunks = []
#         start = 0
#         while start < len(text):
#             end = start + self.chunk_size
#             chunk = text[start:end]
            
#             # Try to break at sentence boundary
#             if end < len(text):
#                 last_period = chunk.rfind(". ")
#                 if last_period > self.chunk_size * 0.5:
#                     chunk = chunk[:last_period + 1]
#                     end = start + last_period + 1

#             if chunk.strip():
#                 chunks.append(chunk)
#             start = end - self.overlap

#         return chunks


# # ─────────────────────────────────────────────────────────────────────────────
# #  STEP 4 — MASTER GENERIC PROCESSOR (Your main entry point)
# # ─────────────────────────────────────────────────────────────────────────────

# class GenericRAGProcessor:
#     """
#     SINGLE CLASS to process ANY file type for RAG indexing.
    
#     Usage:
#         processor = GenericRAGProcessor()
#         result = processor.process("students.csv")
#         result = processor.process("notes.txt")      # auto-detects tabular or free text
#         result = processor.process("report.pdf")
    
#     No hardcoding. No column names. Works with any file.
#     This is the INDUSTRY STANDARD approach.
#     """

#     def __init__(self, chunk_size: int = 500, top_k_chunks: int = 50):
#         self.tabular_processor = TabularProcessor()
#         self.text_processor = FreeTextProcessor(chunk_size=chunk_size)
#         self.top_k_chunks = top_k_chunks

#     def process(self, file_path: str) -> ProcessingResult:
#         """
#         Main method — give it any file path, get back RAG-ready documents.
#         """
#         if not os.path.exists(file_path):
#             raise FileNotFoundError(f"File not found: {file_path}")

#         print(f"\n📂 Processing: {Path(file_path).name}")

#         # ── Step 1: Auto-detect file type ──
#         file_info = FileTypeDetector.detect(file_path)
#         print(f"   Detected type : {file_info['file_type']}")
#         print(f"   Is tabular    : {file_info['is_tabular']}")
#         if file_info.get("delimiter"):
#             delim_name = {",": "comma", "\t": "tab", "|": "pipe", ";": "semicolon"}
#             print(f"   Delimiter     : {delim_name.get(file_info['delimiter'], file_info['delimiter'])}")

#         # ── Step 2: Route to correct processor ──
#         if file_info["is_tabular"]:
#             result = self.tabular_processor.process(file_path, file_info)
#         else:
#             result = self.text_processor.process(file_path, file_info)

#         # ── Step 3: Report ──
#         print(f"   Documents made: {len(result.documents)}")
#         if result.columns_detected:
#             print(f"   Columns found : {', '.join(result.columns_detected)}")
#         if result.warnings:
#             for w in result.warnings:
#                 print(f"   ⚠ Warning     : {w}")

#         return result

#     def process_folder(self, folder_path: str) -> List[ProcessingResult]:
#         """Process all supported files in a folder."""
#         folder = Path(folder_path)
#         supported = {".csv", ".tsv", ".txt", ".md"}
#         results = []
#         for f in folder.iterdir():
#             if f.suffix.lower() in supported:
#                 results.append(self.process(str(f)))
#         return results



# # ─────────────────────────────────────────────────────────────────────────────
# #  STEP 6 — QUERY ENHANCEMENT (Solves your original bug!)
# # ─────────────────────────────────────────────────────────────────────────────

# class QueryEnhancer:
#     """
#     Improves weak queries before sending to the retriever.
#     Solves the "this cv" / "this document" problem generically.
#     """

#     WEAK_PRONOUNS = {"this", "that", "these", "those", "it", "the", "said"}
#     WEAK_PHRASES = [
#         "this document", "this file", "this cv", "this resume",
#         "this data", "this table", "this report", "this record",
#         "the document", "the file", "the data"
#     ]

#     @staticmethod
#     def enhance(query: str, active_file: Optional[str] = None) -> str:
#         """
#         Replace vague references with concrete terms if a file is selected.
        
#         "Please tell me skill of this cv"
#             → "Please tell me skill of the document Syed_Shahzad_Ali.pdf"
        
#         "Who has the highest GPA?"  (with students.csv active)
#             → "Who has the highest GPA in students data?"
#         """
#         enhanced = query

#         if active_file:
#             file_stem = Path(active_file).stem.replace("_", " ").replace("-", " ")
            
#             # Replace weak document references
#             for phrase in QueryEnhancer.WEAK_PHRASES:
#                 if phrase.lower() in enhanced.lower():
#                     enhanced = re.sub(
#                         re.escape(phrase), 
#                         file_stem, 
#                         enhanced, 
#                         flags=re.IGNORECASE
#                     )
#                     break

#         return enhanced

#     @staticmethod
#     def fallback_keyword_search(
#         query: str, 
#         collection,
#         top_k: int = 5,
#         filter_file: Optional[str] = None
#     ) -> List[Dict]:
#         """
#         If semantic search returns 0 results, try keyword-based fallback.
#         Extract nouns/names from query and search directly.
#         """
#         # Extract capitalized words (likely names/entities)
#         keywords = re.findall(r"\b[A-Z][a-z]+\b", query)
        
#         if not keywords:
#             return []

#         # Try each keyword
#         for kw in keywords:
#             where = {"source_file": {"$eq": filter_file}} if filter_file else None
#             results = collection.query(
#                 query_texts=[kw],
#                 n_results=top_k,
#                 where=where,
#                 include=["documents", "metadatas", "distances"]
#             )
#             if results and results["documents"] and results["documents"][0]:
#                 docs = results["documents"][0]
#                 metas = results["metadatas"][0]
#                 dists = results["distances"][0]
#                 return [
#                     {"text": d, "metadata": m, "relevance_score": round(1 - dist, 4)}
#                     for d, m, dist in zip(docs, metas, dists)
#                 ]
#         return []













"""
Document Processing Service
Handles file upload, text extraction, chunking, and vectorization.
Supports: PDF, DOCX, TXT, CSV, TSV — including all table formats.

Classes:
    DocumentProcessor    — Main upload handler (Django integration)
    FileTypeDetector     — Auto-detects file type and delimiter
    TabularProcessor     — Converts CSV/TXT rows → natural language
    FreeTextProcessor    — Sliding-window chunker for plain text
    GenericRAGProcessor  — Master router for any file type
    QueryEnhancer        — Fixes vague queries like "this cv"
"""

# ─────────────────────────────────────────────────────────────────────────────
#  IMPORTS
# ─────────────────────────────────────────────────────────────────────────────
import os
import re
import io
import csv as csv_module
import hashlib
import logging
from pathlib import Path
from dataclasses import dataclass, field
from typing import Dict, Any, List, Optional

import pandas as pd

from django.conf import settings
from django.core.files.uploadedfile import UploadedFile

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────────────────────
#  DATA CLASSES
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class Document:
    """A single indexable unit for ChromaDB / any vector store."""
    text:        str
    metadata:    Dict[str, Any]
    doc_id:      str
    source_file: str
    chunk_index: int           = 0
    row_index:   Optional[int] = None


@dataclass
class ProcessingResult:
    """Result returned after processing any file."""
    documents:        List[Document]
    file_type:        str
    total_rows:       int
    columns_detected: List[str]
    warnings:         List[str] = field(default_factory=list)


# ─────────────────────────────────────────────────────────────────────────────
#  1. DOCUMENT PROCESSOR  — Django upload handler
# ─────────────────────────────────────────────────────────────────────────────

class DocumentProcessor:
    """
    Processes uploaded documents for the RAG system.

    KEY DESIGN — Tabular files (CSV, TSV, tabular TXT):
        Each ROW is stored as its own separate chunk with its own embedding.
        This ensures "find Daniel Clark" retrieves EXACTLY his row,
        not a 30-row blob where he is buried.

    Non-tabular files (PDF, DOCX, plain TXT):
        Standard sliding-window chunking.
    """

    def __init__(
        self,
        vector_store,
        embedding_service,
        chunk_size:    int = None,
        chunk_overlap: int = None,
    ):
        self.vector_store      = vector_store
        self.embedding_service = embedding_service
        self.chunk_size        = chunk_size    or getattr(settings, "CHUNK_SIZE",    800)
        self.chunk_overlap     = chunk_overlap or getattr(settings, "CHUNK_OVERLAP", 100)

    # ─────────────────────────────────────────────────────────────────────────
    #  MAIN ENTRY POINT
    # ─────────────────────────────────────────────────────────────────────────

    async def process_document(
        self,
        file:        UploadedFile,
        document_id: str,
    ) -> Dict[str, Any]:
        """
        Full pipeline: detect type → extract → chunk → embed → store.

        Tabular files  → row-by-row indexing (1 chunk per row)
        Other files    → sliding-window chunking
        """
        name_lower = file.name.lower()
        ct         = (file.content_type or "").lower()

        # ── Detect tabular files ──────────────────────────────────────────
        if (
            name_lower.endswith(".csv") or "csv" in ct or
            name_lower.endswith(".tsv") or
            (name_lower.endswith(".txt") and self._peek_is_tabular(file))
        ):
            logger.info(f"[Processor] TABULAR file: '{file.name}' — indexing row by row")
            return await self._process_tabular(file, document_id)
        else:
            logger.info(f"[Processor] TEXT file: '{file.name}' — sliding-window chunks")
            return await self._process_freetext(file, document_id)

    # ─────────────────────────────────────────────────────────────────────────
    #  TABULAR PIPELINE — one chunk per row
    # ─────────────────────────────────────────────────────────────────────────

    async def _process_tabular(
        self,
        file:        UploadedFile,
        document_id: str,
    ) -> Dict[str, Any]:
        """
        Index CSV / TSV / tabular-TXT files row by row.

        Flow:
            Read file → detect delimiter → parse headers + rows
            → convert each row to "Header: Value | Header: Value ..."
            → embed each row separately → store in ChromaDB

        Result:
            100-row CSV → 100 separate chunks, each with its own embedding
            Query "Daniel Clark" → hits exactly his row → correct answer ✅
        """
        content    = file.read()
        raw        = self._decode(content)
        name_lower = file.name.lower()

        # Detect delimiter
        if name_lower.endswith(".tsv"):
            delimiter = "\t"
        elif name_lower.endswith(".csv"):
            delimiter = ","
        else:
            # TXT — detect automatically from content
            lines     = [l for l in raw.splitlines() if l.strip()]
            delimiter = self._detect_delimiter(lines)

        # Parse rows → list of natural-language strings (one per row)
        row_chunks = self._rows_to_nl_list(raw, delimiter)

        if not row_chunks:
            raise ValueError(f"No data rows found in '{file.name}'")

        logger.info(
            f"[Processor] '{file.name}': {len(row_chunks)} rows "
            f"parsed with delimiter '{delimiter}'"
        )

        # Embed all rows in one batch call
        embeddings = self.embedding_service.embed_texts(row_chunks)

        # Build metadata — IMPORTANT: store both 'source' AND 'document_id'
        # 'source'      → used by document_filter (filename-based filtering)
        # 'document_id' → used by document_id-based filtering
        metadatas = [
            {
                "source":       file.name,
                "document_id":  document_id,
                "content_type": file.content_type or "text/csv",
                "chunk_index":  i,
                "chunk_size":   len(chunk),
                "chunk_type":   "table_row",
                "row_number":   i + 1,
            }
            for i, chunk in enumerate(row_chunks)
        ]

        ids = [f"{document_id}_row_{i}" for i in range(len(row_chunks))]

        self.vector_store.add_documents(
            documents  = row_chunks,
            embeddings = embeddings,
            metadata   = metadatas,
            ids        = ids,
        )

        logger.info(f"[Processor] Stored {len(row_chunks)} row-chunks for '{file.name}'")

        return {
            "chunks_created": len(row_chunks),
            "text_length":    sum(len(r) for r in row_chunks),
            "document_id":    document_id,
            "indexing_mode":  "row_by_row",
        }

    # ─────────────────────────────────────────────────────────────────────────
    #  FREE TEXT PIPELINE — sliding window
    # ─────────────────────────────────────────────────────────────────────────

    async def _process_freetext(
        self,
        file:        UploadedFile,
        document_id: str,
    ) -> Dict[str, Any]:
        """Standard pipeline for PDF, DOCX, plain TXT."""
        text = self._extract_text(file)

        if not text or len(text.strip()) < 10:
            raise ValueError(f"No extractable text found in '{file.name}'")

        logger.info(f"[Processor] Extracted {len(text):,} chars from '{file.name}'")

        chunks = self._chunk_text(text)
        if not chunks:
            raise ValueError(f"No chunks generated from '{file.name}'")

        logger.info(f"[Processor] Created {len(chunks)} sliding-window chunks")

        embeddings = self.embedding_service.embed_texts(chunks)

        metadatas = [
            {
                "source":       file.name,
                "document_id":  document_id,
                "content_type": file.content_type or "application/octet-stream",
                "chunk_index":  i,
                "chunk_size":   len(chunk),
                "chunk_type":   "text_chunk",
            }
            for i, chunk in enumerate(chunks)
        ]

        ids = [f"{document_id}_chunk_{i}" for i in range(len(chunks))]

        self.vector_store.add_documents(
            documents  = chunks,
            embeddings = embeddings,
            metadata   = metadatas,
            ids        = ids,
        )

        logger.info(f"[Processor] Stored {len(chunks)} text-chunks for '{file.name}'")

        return {
            "chunks_created": len(chunks),
            "text_length":    len(text),
            "document_id":    document_id,
            "indexing_mode":  "sliding_window",
        }

    # ─────────────────────────────────────────────────────────────────────────
    #  TABULAR HELPERS
    # ─────────────────────────────────────────────────────────────────────────

    def _peek_is_tabular(self, file: UploadedFile) -> bool:
        """
        Peek at first 2 KB of a TXT file to check if it looks tabular.
        Resets file position after reading so the rest of the pipeline
        can read the file normally.
        """
        try:
            pos    = file.tell()
            sample = file.read(2048)
            file.seek(pos)                      # reset — very important!
            raw    = self._decode(sample)
            lines  = [l.strip() for l in raw.splitlines() if l.strip()][:5]
            if not lines:
                return False
            for delim in (",", "\t", ";", "|"):
                counts = [l.count(delim) for l in lines]
                if len(set(counts)) == 1 and counts[0] >= 1:
                    return True
        except Exception:
            pass
        return False

    @staticmethod
    def _detect_delimiter(lines: List[str]) -> str:
        """
        Return the most consistently used delimiter from sample lines.
        Checks: comma, tab, semicolon, pipe.
        """
        best_delim = ","
        best_count = 0
        for delim in (",", "\t", ";", "|"):
            counts = [l.count(delim) for l in lines[:5]]
            if len(set(counts)) == 1 and counts[0] > best_count:
                best_count = counts[0]
                best_delim = delim
        return best_delim

    def _rows_to_nl_list(self, raw: str, delimiter: str = ",") -> List[str]:
        """
        Parse delimited text and return ONE natural-language string PER ROW.

        This is the CRITICAL method — each row gets its own embedding.

        Input (CSV with 3 data rows):
            Student_ID,First_Name,Last_Name,Age,GPA
            S001,Emma,Johnson,16,3.8
            S024,Daniel,Clark,16,3.5
            S003,Olivia,Brown,17,3.9

        Output (list of 3 strings):
            [
              "Student_ID: S001 | First_Name: Emma | Last_Name: Johnson | Age: 16 | GPA: 3.8",
              "Student_ID: S024 | First_Name: Daniel | Last_Name: Clark | Age: 16 | GPA: 3.5",
              "Student_ID: S003 | First_Name: Olivia | Last_Name: Brown | Age: 17 | GPA: 3.9",
            ]

        When user queries "Daniel Clark":
            → embedding matches row 2 precisely ✅
            → NOT buried in 30-row merged blob ✅
        """
        reader = csv_module.reader(io.StringIO(raw), delimiter=delimiter)
        rows   = [r for r in reader if any(c.strip() for c in r)]

        if not rows:
            return []

        # Single row — no header possible
        if len(rows) == 1:
            return [" | ".join(c.strip() for c in rows[0] if c.strip())]

        headers    = [h.strip() for h in rows[0]]
        row_chunks = []

        for row in rows[1:]:
            pairs = []
            for header, cell in zip(headers, row):
                cell = cell.strip()
                if cell:
                    pairs.append(f"{header}: {cell}")

            if pairs:
                row_chunks.append(" | ".join(pairs))

        logger.info(
            f"[Processor] Parsed {len(row_chunks)} rows "
            f"from {len(headers)} columns (delimiter='{delimiter}')"
        )
        return row_chunks

    # ─────────────────────────────────────────────────────────────────────────
    #  TEXT EXTRACTION ROUTER
    # ─────────────────────────────────────────────────────────────────────────

    def _extract_text(self, file: UploadedFile) -> str:
        """Route to the correct extractor based on file type."""
        content    = file.read()
        name_lower = file.name.lower()
        ct         = (file.content_type or "").lower()

        if name_lower.endswith(".pdf") or "pdf" in ct:
            return self._extract_pdf(content)

        if name_lower.endswith(".docx") or "wordprocessingml" in ct:
            return self._extract_docx(content)

        if name_lower.endswith(".txt") or "text/plain" in ct:
            return self._extract_txt(content)

        raise ValueError(
            f"Unsupported file type: '{file.content_type or file.name}'. "
            "Allowed: PDF, DOCX, TXT, CSV, TSV"
        )

    # ─────────────────────────────────────────────────────────────────────────
    #  PDF
    # ─────────────────────────────────────────────────────────────────────────

    def _extract_pdf(self, content: bytes) -> str:
        """
        Extract text and tables from PDF.
        Prefers pdfplumber (table support); falls back to PyPDF2.
        """
        try:
            import pdfplumber
            parts = []

            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        for raw_table in (page.extract_tables() or []):
                            if raw_table and len(raw_table) >= 2:
                                table_text = self._table_rows_to_text(raw_table)
                                if table_text:
                                    parts.append(table_text)
                    except Exception as exc:
                        logger.warning(f"[Processor] PDF table page {page_num}: {exc}")

                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            parts.append(page_text.strip())
                    except Exception as exc:
                        logger.warning(f"[Processor] PDF text page {page_num}: {exc}")

            result = "\n\n".join(parts).strip()
            if result:
                return result

        except ImportError:
            logger.warning("[Processor] pdfplumber not found — falling back to PyPDF2")

        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            text   = "\n".join(page.extract_text() or "" for page in reader.pages)
            return text.strip()
        except Exception as exc:
            raise ValueError(f"PDF extraction failed: {exc}") from exc

    # ─────────────────────────────────────────────────────────────────────────
    #  DOCX
    # ─────────────────────────────────────────────────────────────────────────

    def _extract_docx(self, content: bytes) -> str:
        """Extract paragraphs and tables from DOCX in document order."""
        try:
            import docx as docx_lib
            doc   = docx_lib.Document(io.BytesIO(content))
            parts = []
            WNS   = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

            for element in doc.element.body:
                tag = element.tag.split("}")[-1]

                if tag == "p":
                    text = "".join(
                        n.text or "" for n in element.iter()
                        if n.tag.endswith("}t")
                    )
                    if text.strip():
                        parts.append(text.strip())

                elif tag == "tbl":
                    rows = []
                    for tr in element.findall(f".//{WNS}tr"):
                        cells = [
                            "".join(
                                n.text or "" for n in tc.iter()
                                if n.tag.endswith("}t")
                            ).strip()
                            for tc in tr.findall(f".//{WNS}tc")
                        ]
                        if cells:
                            rows.append(cells)
                    if rows:
                        table_text = self._table_rows_to_text(rows)
                        if table_text:
                            parts.append(table_text)

            return "\n\n".join(parts).strip()

        except Exception as exc:
            raise ValueError(f"DOCX extraction failed: {exc}") from exc

    # ─────────────────────────────────────────────────────────────────────────
    #  TXT (plain text only — tabular TXT handled by _process_tabular)
    # ─────────────────────────────────────────────────────────────────────────

    def _extract_txt(self, content: bytes) -> str:
        """
        Extract plain text. Called only for non-tabular TXT files.
        Tabular TXT files are handled by _process_tabular() directly.
        """
        raw = self._decode(content)
        return raw.strip()

    # ─────────────────────────────────────────────────────────────────────────
    #  TABLE → NATURAL LANGUAGE (for PDF/DOCX tables)
    # ─────────────────────────────────────────────────────────────────────────

    def _table_rows_to_text(self, rows: List[List]) -> str:
        """
        Convert list-of-lists table (first row = headers) to
        natural-language strings joined by newlines.

        Used for PDF and DOCX embedded tables.
        [['Name', 'Age'], ['Ahmed', '25']] → "Name: Ahmed | Age: 25"
        """
        if not rows or len(rows) < 2:
            return ""

        cleaned = [
            [str(c).strip() if c is not None else "" for c in row]
            for row in rows
        ]

        headers   = cleaned[0]
        data_rows = [
            row for row in cleaned[1:]
            if any(cell for cell in row)
            and not all(set(cell) <= {"-", "=", "_", " ", ""} for cell in row)
        ]

        if not data_rows:
            return ""

        result = []
        for row in data_rows:
            pairs = [
                f"{(headers[i] if i < len(headers) and headers[i] else f'col_{i}')}: {cell}"
                for i, cell in enumerate(row) if cell
            ]
            if pairs:
                result.append(" | ".join(pairs))

        return "\n".join(result)

    # ─────────────────────────────────────────────────────────────────────────
    #  CHUNKING (for non-tabular files)
    # ─────────────────────────────────────────────────────────────────────────

    def _chunk_text(self, text: str) -> List[str]:
        """Word-based sliding-window chunker for plain text."""
        if not text:
            return []

        words = text.split()

        if len(words) <= self.chunk_size:
            return [" ".join(words)]

        chunks = []
        start  = 0

        while start < len(words):
            end   = min(start + self.chunk_size, len(words))
            chunk = " ".join(words[start:end])
            if chunk.strip():
                chunks.append(chunk.strip())
            if end >= len(words):
                break
            start = end - self.chunk_overlap

        return chunks

    # ─────────────────────────────────────────────────────────────────────────
    #  DECODE HELPER
    # ─────────────────────────────────────────────────────────────────────────

    @staticmethod
    def _decode(file_bytes: bytes) -> str:
        """Try common encodings safely; never crash on bad bytes."""
        for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
            try:
                return file_bytes.decode(enc)
            except UnicodeDecodeError:
                continue
        return file_bytes.decode("utf-8", errors="replace")


# ─────────────────────────────────────────────────────────────────────────────
#  2. FILE TYPE DETECTOR
# ─────────────────────────────────────────────────────────────────────────────

class FileTypeDetector:
    """Auto-detects file type, delimiter, and encoding. No hardcoding."""

    @staticmethod
    def detect(file_path: str) -> Dict[str, Any]:
        path = Path(file_path)
        ext  = path.suffix.lower()

        info = {
            "extension":  ext,
            "file_type":  "unknown",
            "is_tabular": False,
            "delimiter":  None,
            "has_header": False,
            "encoding":   FileTypeDetector._detect_encoding(file_path),
        }

        if ext == ".csv":
            info.update(file_type="csv", is_tabular=True, delimiter=",", has_header=True)
        elif ext == ".tsv":
            info.update(file_type="tsv", is_tabular=True, delimiter="\t", has_header=True)
        elif ext == ".txt":
            info.update(FileTypeDetector._analyze_txt(file_path, info["encoding"]))
        elif ext in {".pdf", ".docx", ".doc"}:
            info.update(file_type=ext.lstrip("."), is_tabular=False)

        return info

    @staticmethod
    def _detect_encoding(file_path: str) -> str:
        for enc in ("utf-8", "latin-1", "utf-16", "cp1252"):
            try:
                with open(file_path, "r", encoding=enc) as f:
                    f.read(1024)
                return enc
            except (UnicodeDecodeError, OSError):
                continue
        return "latin-1"

    @staticmethod
    def _analyze_txt(file_path: str, encoding: str) -> Dict[str, Any]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                sample = [f.readline() for _ in range(10)]
            sample = [l.strip() for l in sample if l.strip()]
        except Exception:
            return {"file_type": "txt_freetext", "is_tabular": False}

        if not sample:
            return {"file_type": "txt_freetext", "is_tabular": False}

        counts = {d: sum(l.count(d) for l in sample) for d in (",", "\t", "|", ";")}
        best   = max(counts, key=counts.get)

        if counts[best] / max(len(sample), 1) >= 1:
            field_counts  = [len(l.split(best)) for l in sample]
            is_consistent = len(set(field_counts)) <= 2
            if is_consistent:
                first_fields = sample[0].split(best)
                has_header   = all(
                    not re.match(r"^[0-9.]+$", f.strip())
                    for f in first_fields[:5]
                )
                return {
                    "file_type":  "txt_tabular",
                    "is_tabular": True,
                    "delimiter":  best,
                    "has_header": has_header,
                }

        return {"file_type": "txt_freetext", "is_tabular": False}


# ─────────────────────────────────────────────────────────────────────────────
#  3. TABULAR PROCESSOR
# ─────────────────────────────────────────────────────────────────────────────

class TabularProcessor:
    """
    Converts ANY tabular file into RAG-ready Documents.
    Column names auto-detected — never hardcoded.
    Each row = one Document with its own embedding.
    """

    def process(self, file_path: str, file_info: Dict) -> ProcessingResult:
        try:
            df = self._load_dataframe(file_path, file_info)
        except Exception as exc:
            return ProcessingResult(
                documents=[], file_type=file_info["file_type"],
                total_rows=0, columns_detected=[],
                warnings=[f"Failed to load: {exc}"],
            )

        if df.empty:
            return ProcessingResult(
                documents=[], file_type=file_info["file_type"],
                total_rows=0, columns_detected=[],
                warnings=["File is empty"],
            )

        df.columns = [self._clean_col(c) for c in df.columns]
        df         = df.dropna(how="all")
        col_types  = self._detect_col_types(df)
        filename   = Path(file_path).name
        documents  = []

        for idx, row in df.iterrows():
            documents.append(self._row_to_document(row, idx, col_types, filename))

        if len(df) > 10:
            documents.extend(self._summary_chunks(df, filename, col_types))

        warnings = []
        if len(df) > 1000:
            warnings.append(f"Large file ({len(df)} rows) — consider batch indexing.")

        logger.info(f"[TabularProcessor] '{filename}' → {len(documents)} documents")

        return ProcessingResult(
            documents        = documents,
            file_type        = file_info["file_type"],
            total_rows       = len(df),
            columns_detected = list(df.columns),
            warnings         = warnings,
        )

    def _load_dataframe(self, file_path: str, file_info: Dict) -> "pd.DataFrame":
        enc   = file_info.get("encoding", "utf-8")
        delim = file_info.get("delimiter", ",")
        ext   = file_info.get("extension", "")

        if ext == ".csv":
            try:
                return pd.read_csv(file_path, encoding=enc, dtype=str)
            except Exception:
                return pd.read_csv(file_path, encoding=enc, sep=None, engine="python", dtype=str)

        if file_info["file_type"] == "txt_tabular":
            return pd.read_csv(
                file_path, sep=delim, encoding=enc, dtype=str,
                header=0 if file_info.get("has_header") else None,
                on_bad_lines="skip",
            )

        return pd.DataFrame()

    @staticmethod
    def _clean_col(col: str) -> str:
        col = re.sub(r"[_\-]+", " ", str(col).strip())
        return re.sub(r"\s+", " ", col).strip().title()

    @staticmethod
    def _detect_col_types(df: "pd.DataFrame") -> Dict[str, str]:
        types = {}
        for col in df.columns:
            sample = df[col].dropna().head(10)
            n      = max(len(sample), 1)

            numeric_ratio = sum(
                1 for v in sample if re.match(r"^-?[\d.]+$", str(v).strip())
            ) / n

            if numeric_ratio >= 0.8:
                col_lower = col.lower()
                if any(kw in col_lower for kw in ("id", "code", "no", "number")):
                    types[col] = "identifier"
                elif sum(1 for v in sample if re.match(r"^(19|20)\d{2}$", str(v).strip())) / n >= 0.8:
                    types[col] = "year"
                else:
                    types[col] = "numeric"
            elif sum(1 for v in sample if "@" in str(v)) / n >= 0.5:
                types[col] = "email"
            else:
                types[col] = "text"

        return types

    def _row_to_document(
        self,
        row:       "pd.Series",
        row_index: int,
        col_types: Dict[str, str],
        filename:  str,
    ) -> Document:
        record_type = self._record_type(filename)
        parts       = []
        metadata    = {
            "source_file": filename,
            "row_index":   row_index,
            "record_type": record_type,
        }

        for col in row.index:
            val = row[col]
            if pd.isna(val) or str(val).strip() in ("", "nan", "None"):
                continue
            val = str(val).strip()
            parts.append(f"{col}: {val}")
            metadata[col.lower().replace(" ", "_")] = val

        fields_text = " | ".join(parts)
        name        = self._extract_name(row)
        text        = (
            f"{name} — {record_type} from {filename}:\n{fields_text}"
            if name else
            f"{record_type} from {filename}:\n{fields_text}"
        )

        doc_id = hashlib.md5(f"{filename}_{row_index}".encode()).hexdigest()[:16]

        return Document(
            text=text, metadata=metadata, doc_id=doc_id,
            source_file=filename, row_index=row_index,
        )

    def _summary_chunks(
        self,
        df:        "pd.DataFrame",
        filename:  str,
        col_types: Dict[str, str],
    ) -> List[Document]:
        docs        = []
        record_type = self._record_type(filename)
        col_list    = ", ".join(df.columns)

        overview = (
            f"Overview of {filename}: Contains {len(df)} {record_type}s. "
            f"Columns: {col_list}. "
        )

        for col, ctype in col_types.items():
            if ctype == "numeric":
                try:
                    vals = pd.to_numeric(df[col], errors="coerce").dropna()
                    if len(vals):
                        overview += (
                            f"{col} — min: {vals.min():.2f}, "
                            f"max: {vals.max():.2f}, avg: {vals.mean():.2f}. "
                        )
                except Exception:
                    pass

        docs.append(Document(
            text=overview,
            metadata={"source_file": filename, "chunk_type": "summary"},
            doc_id=hashlib.md5(f"{filename}_overview".encode()).hexdigest()[:16],
            source_file=filename,
        ))

        for col, ctype in col_types.items():
            if ctype == "text" and 1 < df[col].nunique() < 30:
                unique_vals = df[col].dropna().unique().tolist()
                vals_str    = ", ".join(str(v) for v in unique_vals[:25])
                cat_text    = (
                    f"All unique values for '{col}' in {filename}: {vals_str}. "
                    f"Total unique: {len(unique_vals)}."
                )
                docs.append(Document(
                    text=cat_text,
                    metadata={"source_file": filename, "chunk_type": "categorical", "column": col},
                    doc_id=hashlib.md5(f"{filename}_{col}_cat".encode()).hexdigest()[:16],
                    source_file=filename,
                ))

        return docs

    @staticmethod
    def _record_type(filename: str) -> str:
        stem = Path(filename).stem.lower()
        stem = re.sub(r"[_\-\.]", " ", stem)
        stem = re.sub(r"\d+", "", stem)
        stem = re.sub(r"\b(data|file|report|sheet|export|list)\b", "", stem)
        stem = re.sub(r"\s+", " ", stem).strip()
        if not stem:
            return "Record"
        if stem.endswith("s") and len(stem) > 4:
            stem = stem[:-1]
        return stem.title() + " record"

    @staticmethod
    def _extract_name(row: "pd.Series") -> Optional[str]:
        row_lower = {k.lower(): v for k, v in row.items()}

        for key in ("name", "full name", "fullname", "title"):
            val = str(row_lower.get(key, "")).strip()
            if val and val not in ("nan", "None", ""):
                return val

        first = str(row_lower.get("first name", row_lower.get("firstname", ""))).strip()
        last  = str(row_lower.get("last name",  row_lower.get("lastname",  ""))).strip()
        if first and last and first != "nan" and last != "nan":
            return f"{first} {last}"

        return None


# ─────────────────────────────────────────────────────────────────────────────
#  4. FREE TEXT PROCESSOR
# ─────────────────────────────────────────────────────────────────────────────

class FreeTextProcessor:
    """Sliding-window chunker for non-tabular text files."""

    def __init__(self, chunk_size: int = 500, overlap: int = 50):
        self.chunk_size = chunk_size
        self.overlap    = overlap

    def process(self, file_path: str, file_info: Dict) -> ProcessingResult:
        enc      = file_info.get("encoding", "utf-8")
        filename = Path(file_path).name

        try:
            with open(file_path, "r", encoding=enc) as f:
                text = f.read()
        except Exception as exc:
            return ProcessingResult(
                documents=[], file_type="txt_freetext",
                total_rows=0, columns_detected=[],
                warnings=[f"Could not read: {exc}"],
            )

        chunks    = self._split(text)
        documents = []

        for i, chunk in enumerate(chunks):
            doc_id = hashlib.md5(f"{filename}_{i}".encode()).hexdigest()[:16]
            documents.append(Document(
                text=chunk.strip(),
                metadata={"source_file": filename, "chunk_index": i, "chunk_type": "text"},
                doc_id=doc_id,
                source_file=filename,
                chunk_index=i,
            ))

        logger.info(f"[FreeTextProcessor] '{filename}' → {len(documents)} chunks")

        return ProcessingResult(
            documents=documents, file_type="txt_freetext",
            total_rows=len(chunks), columns_detected=[],
        )

    def _split(self, text: str) -> List[str]:
        chunks = []
        start  = 0

        while start < len(text):
            end   = start + self.chunk_size
            chunk = text[start:end]

            if end < len(text):
                boundary = chunk.rfind(". ")
                if boundary > self.chunk_size * 0.5:
                    chunk = chunk[: boundary + 1]
                    end   = start + boundary + 1

            if chunk.strip():
                chunks.append(chunk)

            start = end - self.overlap

        return chunks


# ─────────────────────────────────────────────────────────────────────────────
#  5. GENERIC RAG PROCESSOR
# ─────────────────────────────────────────────────────────────────────────────

class GenericRAGProcessor:
    """Master router — give it any file, get back RAG-ready Documents."""

    def __init__(self, chunk_size: int = 500):
        self.tabular_processor = TabularProcessor()
        self.text_processor    = FreeTextProcessor(chunk_size=chunk_size)

    def process(self, file_path: str) -> ProcessingResult:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_info = FileTypeDetector.detect(file_path)

        logger.info(
            f"[GenericRAGProcessor] '{Path(file_path).name}' | "
            f"type={file_info['file_type']} | tabular={file_info['is_tabular']}"
        )

        if file_info["is_tabular"]:
            return self.tabular_processor.process(file_path, file_info)
        else:
            return self.text_processor.process(file_path, file_info)

    def process_folder(self, folder_path: str) -> List[ProcessingResult]:
        supported = {".csv", ".tsv", ".txt", ".md"}
        results   = []
        for f in Path(folder_path).iterdir():
            if f.suffix.lower() in supported:
                results.append(self.process(str(f)))
        return results


# ─────────────────────────────────────────────────────────────────────────────
#  6. QUERY ENHANCER
# ─────────────────────────────────────────────────────────────────────────────

class QueryEnhancer:
    """
    Fixes vague queries before they reach the retriever.

    "skills of this cv"  +  active_file="Syed_Shahzad_Ali.pdf"
    → "skills of Syed Shahzad Ali"
    """

    WEAK_PHRASES = [
        "this document", "this file",   "this cv",     "this resume",
        "this data",     "this table",  "this report", "this record",
        "the document",  "the file",    "the data",    "the cv",
        "the resume",    "the report",  "this pdf",    "the pdf",
    ]

    @staticmethod
    def enhance(query: str, active_file: Optional[str] = None) -> str:
        if not active_file:
            return query

        label    = Path(active_file).stem.replace("_", " ").replace("-", " ")
        enhanced = query

        for phrase in QueryEnhancer.WEAK_PHRASES:
            if phrase.lower() in enhanced.lower():
                enhanced = re.sub(
                    re.escape(phrase), label, enhanced, flags=re.IGNORECASE,
                )
                break

        if enhanced != query:
            logger.info(f"[QueryEnhancer] '{query}' → '{enhanced}'")

        return enhanced

    @staticmethod
    def fallback_keyword_search(
        query:       str,
        collection,
        top_k:       int           = 5,
        filter_file: Optional[str] = None,
    ) -> List[Dict]:
        keywords = re.findall(r"\b[A-Z][a-z]+\b", query)

        if not keywords:
            logger.warning("[QueryEnhancer] Fallback: no keywords found")
            return []

        logger.info(f"[QueryEnhancer] Fallback keywords: {keywords}")
        where = {"source": {"$eq": filter_file}} if filter_file else None

        for kw in keywords:
            try:
                results = collection.query(
                    query_texts=[kw], n_results=top_k, where=where,
                    include=["documents", "metadatas", "distances"],
                )
                if results and results["documents"] and results["documents"][0]:
                    logger.info(f"[QueryEnhancer] Fallback hit: '{kw}'")
                    return [
                        {
                            "text":            doc,
                            "metadata":        meta,
                            "relevance_score": round(1 - dist, 4),
                        }
                        for doc, meta, dist in zip(
                            results["documents"][0],
                            results["metadatas"][0],
                            results["distances"][0],
                        )
                    ]
            except Exception as exc:
                logger.warning(f"[QueryEnhancer] Fallback error '{kw}': {exc}")
                continue

        logger.warning("[QueryEnhancer] Fallback: no results")
        return []